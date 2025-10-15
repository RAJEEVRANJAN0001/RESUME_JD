"""
Azure Document Intelligence service for parsing PDF/DOCX files.
Includes enhanced fallback to pdfplumber and python-docx.
"""
from azure.core.credentials import AzureKeyCredential
from azure.ai.formrecognizer import DocumentAnalysisClient
from app.config import settings
import PyPDF2
import pdfplumber
import docx
import io
import re
from typing import Optional


class AzureDocumentParser:
    """Parse documents using Azure Document Intelligence with fallback."""
    
    def __init__(self):
        """Initialize Azure client."""
        try:
            self.client = DocumentAnalysisClient(
                endpoint=settings.AZURE_DOC_INTELLIGENCE_ENDPOINT,
                credential=AzureKeyCredential(settings.AZURE_DOC_INTELLIGENCE_KEY)
            )
            self.available = True
        except Exception as e:
            print(f"⚠️ Azure Document Intelligence not available: {e}")
            self.available = False
    
    async def parse_document(self, file_content: bytes, filename: str) -> str:
        """
        Parse document and extract text.
        Tries Azure first, then falls back to local parsers.
        """
        # Try Azure Document Intelligence
        if self.available:
            try:
                text = await self._parse_with_azure(file_content)
                if text:
                    return text
            except Exception as e:
                print(f"⚠️ Azure parsing failed: {e}, falling back...")
        
        # Fallback to local parsers
        if filename.lower().endswith('.pdf'):
            return self._parse_pdf(file_content)
        elif filename.lower().endswith('.docx'):
            return self._parse_docx(file_content)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")
    
    async def _parse_with_azure(self, file_content: bytes) -> Optional[str]:
        """Parse using Azure Document Intelligence."""
        poller = self.client.begin_analyze_document(
            "prebuilt-document",
            document=file_content
        )
        result = poller.result()
        
        # Extract text content
        text_content = []
        for page in result.pages:
            for line in page.lines:
                text_content.append(line.content)
        
        return "\n".join(text_content)
    
    def _parse_pdf(self, file_content: bytes) -> str:
        """Enhanced PDF parser using pdfplumber first, then PyPDF2 fallback."""
        # Try pdfplumber first (better text extraction)
        try:
            pdf_file = io.BytesIO(file_content)
            text_content = []
            
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    # Extract text with layout preservation
                    text = page.extract_text(layout=True)
                    if text:
                        text_content.append(text)
            
            if text_content:
                full_text = "\n".join(text_content)
                cleaned_text = self._clean_text(full_text)
                if len(cleaned_text) > 50:  # Ensure we got meaningful content
                    print(f"✅ PDF parsed successfully with pdfplumber ({len(cleaned_text)} chars)")
                    return cleaned_text
        except Exception as e:
            print(f"⚠️  pdfplumber failed: {e}, trying PyPDF2...")
        
        # Fallback to PyPDF2
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
            
            full_text = "\n".join(text_content)
            cleaned_text = self._clean_text(full_text)
            if cleaned_text.strip():
                print(f"✅ PDF parsed with PyPDF2 fallback ({len(cleaned_text)} chars)")
                return cleaned_text
        except Exception as e:
            print(f"❌ PyPDF2 also failed: {e}")
        
        return "Unable to extract text from PDF. Please try a different file format."
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text to improve parsing quality."""
        # Fix common spacing issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between camelCase
        text = re.sub(r'(\d)([A-Z])', r'\1 \2', text)  # Add space between number and letter
        text = re.sub(r'([a-z])(\d)', r'\1 \2', text)  # Add space between letter and number
        
        # Fix bullet points
        text = re.sub(r'•', '\n• ', text)
        text = re.sub(r'●', '\n• ', text)
        text = re.sub(r'○', '\n• ', text)
        text = re.sub(r'▪', '\n• ', text)
        text = re.sub(r'►', '\n• ', text)
        
        # Fix multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Fix multiple newlines
        text = re.sub(r'\n\n+', '\n\n', text)
        
        return text.strip()
    
    def _parse_docx(self, file_content: bytes) -> str:
        """Enhanced DOCX parser using python-docx with better structure preservation."""
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        
        text_content = []
        
        # Extract paragraphs with formatting awareness
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Detect headings by style or bold text
                if paragraph.style.name.startswith('Heading'):
                    text_content.append(f"\n{text.upper()}\n")
                else:
                    text_content.append(text)
        
        # Extract tables (often used in resumes)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        full_text = "\n".join(text_content)
        return self._clean_text(full_text) if full_text.strip() else "Unable to extract text from DOCX"


# Global parser instance
azure_parser = AzureDocumentParser()
